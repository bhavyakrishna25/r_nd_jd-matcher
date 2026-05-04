from io import BytesIO

from docx import Document
from PyPDF2 import PdfReader


class ResumeParserError(Exception):
    pass


def get_supported_extension(filename: str) -> str:
    lower_name = filename.lower().strip()

    if lower_name.endswith(".pdf"):
        return ".pdf"
    if lower_name.endswith(".docx"):
        return ".docx"

    raise ResumeParserError("Only PDF and DOCX files are supported.")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    pdf_stream = BytesIO(file_bytes)
    reader = PdfReader(pdf_stream)

    pages_text = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages_text.append(page_text)

    return "\n".join(pages_text).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    docx_stream = BytesIO(file_bytes)
    document = Document(docx_stream)

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    table_rows = []
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_rows.append(row_text)

    return "\n".join(paragraphs + table_rows).strip()


def extract_text_from_resume(filename: str, file_bytes: bytes) -> str:
    extension = get_supported_extension(filename)

    if extension == ".pdf":
        extracted_text = extract_text_from_pdf(file_bytes)
    else:
        extracted_text = extract_text_from_docx(file_bytes)

    if not extracted_text.strip():
        raise ResumeParserError("No readable text was found in the uploaded resume.")

    return extracted_text
